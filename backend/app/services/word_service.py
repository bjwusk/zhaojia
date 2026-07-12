"""
Word 造价报告填充服务（v3 - 完整填充项目背景/材料清单/摘要/ZIP导出）
完整的费用链：人工费→措施费→直接工程费→管理费→利润→税金→预备费→总造价
"""
import os, re, shutil, zipfile
from datetime import datetime

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
TEMPLATE_PATH = os.path.join(BASE_DIR, "第三方软件造价评估报告模板_（BSCEA提供）_2024版.docx")
EXPORT_DIR = os.path.join(BASE_DIR, "backend", "exports")

PDR_MEDIAN = 6.72

# 措施费率
SURVEY_FEE_RATE = 0.05
TEST_FEE_RATE = 0.08
TRAINING_FEE_RATE = 0.03
THIRD_PARTY_TEST_RATE = 0.04
SECURITY_TEST_RATE = 0.03
MANAGEMENT_RATE = 0.10
PROFIT_RATE = 0.08
TAX_RATE_NORMAL = 0.06
TAX_RATE_SIMPLE = 0.03
BASIC_RESERVE_RATE = 0.05
RISK_RATE = 0.03


def calc_full_cost(labor_cost: float, params: dict = None) -> dict:
    """完整费用链计算"""
    if params is None:
        params = {}
    tax_type = params.get("tax_type", "一般计税")
    measure_items = params.get("measure_items", ["survey_fee", "test_fee", "training"])
    include_mgmt = params.get("include_management", True)
    include_profit = params.get("include_profit", True)
    include_reserve = params.get("include_basic_reserve", True)
    include_risk = params.get("include_risk", False)

    direct_labor = round(labor_cost, 4)

    measure_amount = 0.0
    measure_details = []
    rate_map = {"survey_fee": ("需求调研费", SURVEY_FEE_RATE),
                "test_fee": ("测试费", TEST_FEE_RATE),
                "training": ("培训费", TRAINING_FEE_RATE),
                "third_party_test": ("第三方测评", THIRD_PARTY_TEST_RATE),
                "security_test": ("安全测评", SECURITY_TEST_RATE)}
    for key in measure_items:
        if key in rate_map:
            name, rate = rate_map[key]
            v = round(direct_labor * rate, 4)
            measure_details.append({"name": name, "rate": rate, "amount": v})
            measure_amount += v

    measure_total = round(measure_amount, 4)
    engineering_cost = round(direct_labor + measure_total, 4)
    management_fee = round(engineering_cost * MANAGEMENT_RATE, 4) if include_mgmt else 0
    profit = round(engineering_cost * PROFIT_RATE, 4) if include_profit else 0
    tax_rate = TAX_RATE_NORMAL if tax_type == "一般计税" else TAX_RATE_SIMPLE
    tax_base = engineering_cost + management_fee + profit
    tax = round(tax_base * tax_rate, 4)
    reserve_base = tax_base + tax
    basic_reserve = round(reserve_base * BASIC_RESERVE_RATE, 4) if include_reserve else 0
    risk_fee = round(reserve_base * RISK_RATE, 4) if include_risk else 0
    total = round(engineering_cost + management_fee + profit + tax + basic_reserve + risk_fee, 4)

    return {
        "direct_labor": direct_labor,
        "measure_details": measure_details,
        "measure_total": measure_total,
        "engineering_cost": engineering_cost,
        "management_fee": management_fee,
        "profit": profit,
        "tax_type": tax_type,
        "tax_rate": tax_rate,
        "tax": tax,
        "basic_reserve": basic_reserve,
        "risk_fee": risk_fee,
        "total": total,
    }


def fill_report(project_name: str, result: dict, fp_items: list, params: dict,
                doc_text: str = "", doc_filename: str = "", output_path: str = None) -> str:
    """将测算结果填入Word报告模板"""
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Word模板不存在: {TEMPLATE_PATH}")

    from docx import Document

    doc = Document(TEMPLATE_PATH)
    os.makedirs(EXPORT_DIR, exist_ok=True)

    labor_cost = result.get("labor_cost_median", 0) or 0
    fee_params = params.get("fee_params", {}) if params else {}
    cost_chain = calc_full_cost(labor_cost, fee_params)

    # ── 构建项目背景（从文档中提取前2000字） ───────────
    background_text = ""
    if doc_text and len(doc_text.strip()) > 50:
        # 取文档前2000字作为项目背景摘要
        background_text = doc_text.strip()[:2000]
        # 去掉可能不完整的最后一句话
        last_punc = max(background_text.rfind("。"), background_text.rfind("；"), background_text.rfind("；"))
        if last_punc > 100:
            background_text = background_text[:last_punc+1]
        background_text += "……(项目完整建设方案详见原文文档)"

    if not background_text:
        background_text = f"本项目为{project_name or '软件项目'}，主要建设内容包括需求分析、系统设计、开发实施、测试验收等阶段。本次造价评估基于项目需求文档进行功能点分析，采用BSCEA标准方法进行测算。"

    # ── 替换段落占位符 ──────────────────────────
    paragraphs = doc.paragraphs

    # 构建替换映射
    unit_price = result.get("unit_price", 3.2198)
    city = params.get("city", "北京") if params else "北京"
    fp_method_label = "预估功能点法" if params.get("fp_method") == "estimated" else "估算功能点法"

    replacements = {
        "XXXXXXXXXXXXXXXXXXX项目": (project_name or "软件项目") + "项目",
        "项目名称：XXXXXX": f"项目名称：{project_name or '待定'}",
        "委托单位：XXXXXX": f"委托单位：{params.get('client', '建设单位') if params else '建设单位'}",
        "评估内容：XXXXXX": f"评估内容：{project_name or '待定'}造价评估",
        "评估机构：XXXXXX": "评估机构：软件造价自动生成器",
        "评估专家：XXX": f"评估专家：{params.get('creator', '系统生成') if params else '系统生成'}",
        "证书编号：XXXXXX": f"证书编号：{params.get('cert_no', '') if params else ''}",
        "审核专家：XXX": f"审核专家：{params.get('reviewer', '') if params else ''}",
        "XXXX年XX月XX日": datetime.now().strftime("%Y年%m月%d日"),

        # 摘要部分
        "XXXXXXXXXX受 XXXXXXXXXX委托": f"{'建设单位' if params else '我方'}受 {'委托单位' if params else '贵单位'}委托",
        # skip problematic key
        "项目功能点规模为XXX FP": f"项目功能点规模为{result.get('adjusted_fp', 0) or 0} FP",
        "该项目软件开发工作量上限值为XXX人天": f"该项目软件开发工作量上限值为{result.get('workload_upper', 0) or 0}人天",
        "中值为XXX人天": f"中值为{result.get('workload_median', 0) or 0}人天",
        "下限值为XXX人天": f"下限值为{result.get('workload_lower', 0) or 0}人天",
        "该项目软件开发成本上限值为XXX万元": f"该项目软件开发成本上限值为{round(result.get('labor_cost_upper', 0) or 0, 2)}万元",
        "中值为XXX万元": f"中值为{round(result.get('labor_cost_median', 0) or 0, 2)}万元",
        "下限值为XXX万元": f"下限值为{round(result.get('labor_cost_lower', 0) or 0, 2)}万元",

        # 全文
        "XXXX公司:": f"{params.get('client', '建设单位') if params else '建设单位'}:",
        "对XXXX的功能点规模": f"对{project_name or '本项目'}的功能点规模",

    }

    for p in paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in p.text:
                for run in p.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)

    # 再次遍历段落替换剩余占位符
    for p in paragraphs:
        text = p.text.strip()
        if "??????????" in text:
            for run in p.runs:
                if "XXXXXXXXXX" in run.text:
                    run.text = run.text.replace("XXXXXXXXXX", project_name or "???", 1)
                if "XXXXXXXXXXX" in run.text:
                    run.text = run.text.replace("XXXXXXXXXXX", project_name or "???", 1)
                if "XXXX" in run.text and len(run.text.strip().replace("X","").strip()) < 5:
                    run.text = run.text.replace("XXXX", project_name or "本项目", 1)
            break

    # ── 填充项目背景段落（p[75]通常） ──────────────
    for p in paragraphs:
        if p.text.strip() == "项目背景":
            # 下一段就是项目背景内容
            pass
        # 找纯XXX占位符段落（背景内容）
        t = p.text.strip()
        if t and all(c == "X" or c in "，。、" for c in t) and len(t) > 20:
            for run in p.runs:
                if run.text.strip() and all(c == "X" or c in "，。、" for c in run.text.strip()):
                    run.text = background_text
                    break

    # ── 填充评估材料清单（文件名） ──────────────────
    material_items = []
    if doc_filename:
        material_items.append(doc_filename)
    else:
        material_items.append(f"{project_name or '软件项目'}需求文档")

    # 找"评估材料清单"下面的列表段落
    found_material_header = False
    for p in paragraphs:
        t = p.text.strip()
        if t == "评估材料清单":
            found_material_header = True
            continue
        if found_material_header:
            if t == "XXXXX":
                for run in p.runs:
                    if "XXXXX" in run.text:
                        run.text = run.text.replace("XXXXX", material_items[0] if material_items else "需求文档", 1)
                    if "XXXXX" in run.text:
                        run.text = run.text.replace("XXXXX", "软件造价评估委托书", 1)
                    if "XXXXX" in run.text:
                        run.text = run.text.replace("XXXXX", "系统设计方案", 1)
                break
            elif t == "" or t == "……":
                continue
            elif t.startswith("……"):
                continue

    # ── 评估方法段落 ────────────────────────────
    for p in paragraphs:
        if "功能点的计数规则，使用ISO/IEC 24570功能点标准中的XXXXXX方法" in p.text:
            for run in p.runs:
                if "XXXXXX方法" in run.text:
                    run.text = run.text.replace("XXXXXX方法", f"{fp_method_label}方法")
                    break

    # ── 填写结果表格（智能匹配模板行标签） ──────
    if doc.tables:
        table0 = doc.tables[0]
        row_data = [
            ("规模估算结果", str(result.get("adjusted_fp", 0) or 0), "功能点"),
            ("规模变更调整因子", str(result.get("scale_factor", 1) or 1), f"估算时机: {params.get('scale_timing', '交付后') if params else '交付后'}"),
            ("调整后规模", str(result.get("adjusted_fp", 0) or 0), "功能点"),
            ("基准生产率", str(round(result.get("pdr_median", PDR_MEDIAN) or PDR_MEDIAN, 2)), "人时/功能点"),
            ("未调整工作量（中值）", str(round(result.get("workload_median", 0) or 0, 2)), "人天"),
            ("工作量调整因子乘积", str(round(result.get("adj_factor_product", 1.0) or 1.0, 4)), f"应用类型×非功能×完整性×语言×团队"),
            ("调整后工作量（中值）", str(round(result.get("adj_workload_median", 0) or 0, 2)), "人天"),
            ("人月数（中值）", str(round(result.get("person_months_median", 0) or 0, 2)), "人月"),
            ("人工费（中值）", str(round(cost_chain["direct_labor"], 2)), "万元"),
            ("措施费", str(round(cost_chain["measure_total"], 2)), "万元"),
            ("直接工程费", str(round(cost_chain["engineering_cost"], 2)), "万元"),
            ("管理费", str(round(cost_chain["management_fee"], 2)), "万元"),
            ("利润", str(round(cost_chain["profit"], 2)), "万元"),
            ("税金", str(round(cost_chain["tax"], 2)), "万元"),
            ("预备费", str(round(cost_chain["basic_reserve"], 2)), "万元"),
            ("总造价", str(round(cost_chain["total"], 2)), "万元"),
        ]
        # 根据模板行标签智能匹配填充
        row_map = {}
        for ri, row in enumerate(table0.rows):
            label = row.cells[0].text.strip() if row.cells else ''
            if len(row.cells) > 1:
                label2 = row.cells[1].text.strip()
                if label2 and label2 != label:
                    label = label2
            for kw in [label, label.replace(chr(10),' ').replace(chr(13),' ')]:
                for key in row_map:
                    pass
            row_map[ri] = label.replace(chr(10),' ').replace(chr(13),' ').strip()
        for label, val, unit in row_data:
            best_ri = -1
            for ri, lbl in row_map.items():
                if label in lbl or lbl.startswith(label[:4]) or label[:4] in lbl:
                    best_ri = ri
                    break
            if best_ri >= 0:
                row = table0.rows[best_ri]
                if len(row.cells) > 2:
                    set_cell_text(row.cells[2], val)
                if len(row.cells) > 3 and unit:
                    set_cell_text(row.cells[3], unit)

    # ── 功能点清单表 ─────────────────────────────
    if len(doc.tables) > 1:
        table1 = doc.tables[1]
        while len(table1.rows) > 1:
            try:
                table1._tbl.remove(table1.rows[1]._tr)
            except:
                break
        for idx, fp in enumerate(fp_items):
            row = table1.add_row()
            cells = row.cells
            set_cell_text(cells[0], str(idx + 1))
            set_cell_text(cells[1], fp.get("subsystem", ""))
            set_cell_text(cells[2], fp.get("module_l1", ""))
            set_cell_text(cells[3], "")
            set_cell_text(cells[4], fp.get("description", "")[:60])
            set_cell_text(cells[5], fp.get("fp_name", "")[:40])
            set_cell_text(cells[6], fp.get("category", ""))
            set_cell_text(cells[7], str(fp.get("ufp", 0)))
            set_cell_text(cells[8], str(fp.get("reuse_level", "低")))
            set_cell_text(cells[9], str(fp.get("modify_type", "新建")))
            set_cell_text(cells[10], str(fp.get("us", fp.get("ufp", 0))))

    # ── 保存 ────────────────────────────────────
    if not output_path:
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", project_name or "造价报告")
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(EXPORT_DIR, f"{safe_name}_造价评估报告_{ts}.docx")

    doc.save(output_path)
    return output_path


def set_cell_text(cell, text):
    """设置单元格文本"""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text) if text is not None else "")
    return run


def create_zip_package(project_name: str, excel_path: str, word_path: str) -> str:
    """将Excel和Word打包成ZIP文件"""
    safe_name = re.sub(r'[\\/:*?"<>|]', "_", project_name or "造价成果")
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_path = os.path.join(EXPORT_DIR, f"{safe_name}_造价成果_{ts}.zip")

    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        if excel_path and os.path.exists(excel_path):
            zf.write(excel_path, os.path.basename(excel_path))
        if word_path and os.path.exists(word_path):
            zf.write(word_path, os.path.basename(word_path))

    return zip_path
