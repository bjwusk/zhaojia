import os
import json
from datetime import datetime
import xlsxwriter
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.utils.data_loader import load_dev_library, load_csbmk
from app.models.estimate_result import EstimateVersion, CostDetail
from app.models.fp_item import FpItem
from app.models.project import Project

EXPORT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), 'exports')


class ExportService:

    @staticmethod
    def ensure_export_dir():
        os.makedirs(EXPORT_DIR, exist_ok=True)
        return EXPORT_DIR

    @staticmethod
    def export_excel(version_id):
        data = EstimateService._get_export_data(version_id)
        if not data:
            return None
        export_dir = ExportService.ensure_export_dir()
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        pn = data["project_name"]
        filename = f"{pn}_ce_suan_di_gao_{ts}.xlsx"
        filepath = os.path.join(export_dir, filename)

        workbook = xlsxwriter.Workbook(filepath)
        bold = workbook.add_format({"bold": True, "font_size": 12, "bg_color": "#D9E1F2"})
        header_fmt = workbook.add_format({"bold": True, "bg_color": "#4472C4", "font_color": "white", "border": 1})
        cell_fmt = workbook.add_format({"border": 1, "font_size": 10})
        number_fmt = workbook.add_format({"border": 1, "num_format": "#,##0.00", "font_size": 10})
        title_fmt = workbook.add_format({"bold": True, "font_size": 14, "align": "center"})

        # Sheet 1: 项目特征
        ws1 = workbook.add_worksheet("project_info")
        ws1.merge_range("A1:D1", "Project Info", title_fmt)
        headers1 = ["Field", "Value", "Field", "Value"]
        for i, h in enumerate(headers1):
            ws1.write(1, i, h, header_fmt)
        proj = data["project"]
        fields = [
            ("name", proj.name, "client", proj.client),
            ("industry", proj.industry, "stage", proj.stage),
            ("build_mode", proj.build_mode, "region", proj.region),
            ("xinchuang", "Y" if proj.is_xinchuang else "N", "confidential", "Y" if proj.is_confidential else "N"),
            ("estimate_type", "dev" if proj.estimate_type == "dev" else "ops", "creator", proj.creator),
        ]
        for r, (a, b, c, d) in enumerate(fields):
            ws1.write(r+2, 0, a, cell_fmt)
            ws1.write(r+2, 1, str(b) if b else "", cell_fmt)
            ws1.write(r+2, 2, c, cell_fmt)
            ws1.write(r+2, 3, str(d) if d else "", cell_fmt)
        ws1.set_column("A:A", 18)
        ws1.set_column("B:B", 30)
        ws1.set_column("C:C", 18)
        ws1.set_column("D:D", 30)

        # Sheet 2: 功能点清单
        ws2 = workbook.add_worksheet("fp_items")
        fp_headers = ["seq","subsys","module1","module2","module3","module4","desc","fp_name","cat","UFP","reuse","mod_type","US","note"]
        for i, h in enumerate(fp_headers):
            ws2.write(0, i, h, header_fmt)
        items = data["items"]
        for r, item in enumerate(items):
            for c, key in enumerate(["seq","subsystem","module_l1","module_l2","module_l3","module_l4","description","fp_name","category","ufp","reuse_level","modify_type","us","note"]):
                val = item.get(key, "")
                if isinstance(val, (int, float)):
                    ws2.write(r+1, c, val, number_fmt)
                else:
                    ws2.write(r+1, c, str(val) if val else "", cell_fmt)
        ws2.set_column("A:A", 6)
        ws2.set_column("B:F", 12)
        ws2.set_column("G:H", 20)
        ws2.set_column("I:I", 8)
        ws2.set_column("J:K", 10)
        ws2.set_column("L:L", 10)
        ws2.set_column("M:M", 8)
        ws2.set_column("N:N", 12)

        # Sheet 3: 费用明细
        ws3 = workbook.add_worksheet("cost_details")
        cd_headers = ["category","item","base_amount","rate","amount","note"]
        for i, h in enumerate(cd_headers):
            ws3.write(0, i, h, header_fmt)
        details = data["cost_details"]
        tr = len(details) + 1
        for r, d in enumerate(details):
            ws3.write(r+1, 0, d.get("category", ""), cell_fmt)
            ws3.write(r+1, 1, d.get("item_name", ""), cell_fmt)
            ws3.write(r+1, 2, d.get("base_amount", 0), number_fmt)
            ws3.write(r+1, 3, d.get("rate", 0), number_fmt)
            ws3.write(r+1, 4, d.get("amount", 0), number_fmt)
            ws3.write(r+1, 5, d.get("note", ""), cell_fmt)
        ws3.write(tr, 0, "TOTAL COST", bold)
        ws3.write(tr, 4, data["version"].get("total_cost", 0), bold)
        ws3.set_column("A:A", 14)
        ws3.set_column("B:B", 20)
        ws3.set_column("C:E", 14)
        ws3.set_column("F:F", 20)

        workbook.close()
        return filepath

    @staticmethod
    def export_word(version_id):
        data = EstimateService._get_export_data(version_id)
        if not data:
            return None
        export_dir = ExportService.ensure_export_dir()
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        pn = data["project_name"]
        filename = f"{pn}_zhi_jia_report_{ts}.docx"
        filepath = os.path.join(export_dir, filename)

        doc = Document()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\n\n\n" + data["project_name"])
        run.font.size = Pt(22)
        run.bold = True

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("\nSOFTWARE COST REPORT\n\n")
        run2.font.size = Pt(18)

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        proj = data["project"]
        rd = datetime.now().strftime("%Y-%m-%d")
        info = f"Evaluator: {proj.creator or "___"}\nReviewer: {proj.reviewer or "___"}\nDate: {rd}"
        run3 = p3.add_run(info)
        run3.font.size = Pt(12)
        doc.add_page_break()

        # Summary
        doc.add_heading("Summary", level=1)
        v = data["version"]
        fp_total = v.get("adjusted_fp", 0) or 0
        total = v.get("total_cost", 0) or 0
        pm = v.get("person_months", 0) or 0
        doc.add_paragraph(f"FP: {fp_total}")
        doc.add_paragraph(f"Person-months: {pm}")
        doc.add_paragraph(f"Total cost: {total} (10K CNY)")

        # Detail table
        doc.add_heading("Detail Results", level=1)
        table = doc.add_table(rows=8, cols=2, style="Light Grid Accent 1")
        rd2 = [
            ("FP", str(fp_total)),
            ("Adjusted FP", str(fp_total)),
            ("Workload (days)", str(v.get("adjusted_workload", 0))),
            ("Person-months", str(pm)),
            ("Rate (10K/pm)", str(v.get("person_month_rate", 0))),
            ("Labor cost (10K)", str(v.get("labor_cost", 0))),
            ("Total cost (10K)", str(total)),
        ]
        for i, (k, vl) in enumerate(rd2):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[1].text = vl

        doc.save(filepath)
        return filepath


class EstimateService:
    @staticmethod
    def _get_export_data(version_id):
        version = EstimateVersion.query.get(version_id)
        if not version:
            return None
        project = Project.query.get(version.project_id)
        items = FpItem.query.filter_by(version_id=version_id).order_by(FpItem.seq).all()
        details = CostDetail.query.filter_by(version_id=version_id).all()
        return {
            "project": project,
            "project_name": project.name if project else "Unknown",
            "version": version.to_dict(),
            "items": [item.to_dict() for item in items],
            "cost_details": [d.to_dict() for d in details],
        }
